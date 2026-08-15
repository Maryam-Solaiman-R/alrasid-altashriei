import os, html, io, re, time
from fastapi import FastAPI, HTTPException
from fastapi.responses import HTMLResponse, StreamingResponse
from pydantic import BaseModel
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from app.search_engine import discover

app=FastAPI(title="الراصد التشريعي")

SOURCES=[
("boe","هيئة الخبراء بمجلس الوزراء",["laws.boe.gov.sa","boe.gov.sa"]),
("ncar","المركز الوطني للوثائق والمحفوظات",["ncar.gov.sa"]),
("uqn","جريدة أم القرى",["uqn.gov.sa"]),
("mof","وزارة المالية",["mof.gov.sa"]),
("hrsd","وزارة الموارد البشرية والتنمية الاجتماعية",["hrsd.gov.sa"]),
("lcgpa","هيئة المحتوى المحلي والمشتريات الحكومية",["lcgpa.gov.sa"]),
("expro","هيئة كفاءة الإنفاق والمشروعات الحكومية",["expro.gov.sa"]),
("dga","هيئة الحكومة الرقمية",["dga.gov.sa"]),
("nca","الهيئة الوطنية للأمن السيبراني",["nca.gov.sa"]),
("sdaia","الهيئة السعودية للبيانات والذكاء الاصطناعي (سدايا)",["sdaia.gov.sa"]),
("gca","الديوان العام للمحاسبة",["gca.gov.sa"]),
]
class Ask(BaseModel):
    question:str
    sources:list[str]=["all"]

class ExportReq(Ask):
    answer:str
    citations:list[dict]=[]
    source_names:list[str]=[]

def selected(ids):
    if not ids or "all" in ids: return SOURCES
    wanted=set(ids); return [s for s in SOURCES if s[0] in wanted]

_CACHE={}

def _keywords(q):
    stop={'ما','هي','هو','في','من','على','عن','إلى','الى','تم','التي','الذي','هل','وما','بعد','قبل','مع','هذه','هذا'}
    return [x for x in re.findall(r'[\u0600-\u06FF\d]{2,}',q) if x not in stop][:12]

def _score(item, words):
    txt=(item.get('title','')+' '+item.get('snippet','')+' '+item.get('content','')).lower()
    return sum(3 if w in item.get('title','').lower() else 1 for w in words if w.lower() in txt)

def _extract_facts(text):
    text=re.sub(r'\s+',' ',text or '')
    arts=[]
    for m in re.finditer(r'(?:المادة|مادة)\s*(?:رقم\s*)?[\(（]?([0-9٠-٩]{1,4})[\)）]?',text):
        v=m.group(1)
        if v not in arts: arts.append(v)
    decisions=[]
    for m in re.finditer(r'(?:قرار[^،.]{0,45}?رقم|بالقرار[^،.]{0,30}?رقم)\s*[\(（]?([0-9٠-٩]{1,6})[\)）]?',text):
        v=m.group(1)
        if v not in decisions: decisions.append(v)
    dates=[]
    for pat in [r'\b[0-9٠-٩]{1,2}/[0-9٠-٩]{1,2}/[0-9٠-٩]{2,4}\s*هـ?',r'\b[0-9٠-٩]{1,2}/[0-9٠-٩]{1,2}/[0-9٠-٩]{4}\s*م?']:
        for m in re.finditer(pat,text):
            v=m.group(0)
            if v not in dates: dates.append(v)
    return arts[:20],decisions[:10],dates[:10]

async def _free_search(req):
    ss=selected(req.sources); domains=[d for x in ss for d in x[2]]
    key=(req.question,tuple(domains)); now=time.time()
    if key in _CACHE and now-_CACHE[key][0] < 1800:
        return _CACHE[key][1]
    # Free public discovery: no OpenAI API call.
    results=await discover(req.question,domains,limit=10)
    words=_keywords(req.question)
    results=sorted(results,key=lambda x:_score(x,words),reverse=True)
    citations=[]; alltext=''
    for x in results:
        if not x.get('url'): continue
        citations.append({'title':x.get('title') or 'مصدر رسمي','url':x['url']})
        alltext+=' '+x.get('title','')+' '+x.get('snippet','')+' '+x.get('content','')
    arts,decs,dates=_extract_facts(alltext)
    lines=['## نتائج الراصد','تم البحث في المصادر الرسمية المحددة وتحليل النتائج المتاحة وفقًا لاستعلامك.']
    if arts: lines += ['', '**أرقام المواد التي ظهرت في النتائج:** '+ '، '.join(arts)]
    if decs: lines += ['**أرقام القرارات التي ظهرت:** '+ '، '.join(decs)]
    if dates: lines += ['**التواريخ التي ظهرت:** '+ '، '.join(dates)]
    if results:
        lines += ['', '## أبرز النتائج الرسمية']
        for i,x in enumerate(results[:8],1):
            content=re.sub(r'\s+',' ',x.get('content','') or x.get('snippet','')).strip()
            # Presentation-only cleanup: suppress web-template / JavaScript fragments.
            content=re.sub(r'\{\{.*?\}\}|function\s*\([^)]*\)\s*\{.*?\}|toLocaleTimeString\([^)]*\)|limitBodyText\([^)]*\)', ' ', content, flags=re.I)
            content=re.sub(r'\b(?:Facebook|WhatsApp|Read more|Scroll|Title)\s*:?', ' ', content, flags=re.I)
            content=re.sub(r'\s+',' ',content).strip(' -–—|')
            sn=content[:420]
            for w in words:
                pos=content.find(w)
                if pos>=0:
                    sn=content[max(0,pos-100):pos+430]; break
            if len(sn)>520: sn=sn[:517].rstrip()+'...'
            lines.append(f'### {i}. {x.get("title") or "مصدر رسمي"}')
            meta=[]
            if x.get('document_type'): meta.append(f'نوع المصدر: {x.get("document_type")}')
            try:
                from urllib.parse import urlparse
                host=urlparse(x.get('url','')).netloc.replace('www.','')
                if host: meta.append(f'الجهة/النطاق: {host}')
            except Exception:
                pass
            if meta: lines.append(' | '.join(meta))
            if sn: lines.append(sn)
    else:
        lines += ['', 'لم يتم العثور على نتيجة رسمية مطابقة للاستعلام. جرّب كتابة اسم النظام أو اللائحة أو رقم المادة أو القرار بصورة أكثر تحديدًا.']
    discovery_mode=(results[0].get('discovery') if results else 'none')
    data={'answer':'\n'.join(lines),'citations':citations,'sources':[x[1] for x in ss], 'mode':discovery_mode}
    _CACHE[key]=(now,data)
    return data

@app.post('/analyze')
async def analyze(req:Ask):
    try: return await _free_search(req)
    except Exception as e: raise HTTPException(500,str(e))

def _rtl_paragraph(p):
    # True Word RTL: paragraph bidi + complex-script RTL on every run.
    p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    pPr=p._p.get_or_add_pPr()
    bidi=pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi=OxmlElement('w:bidi'); pPr.append(bidi)
    bidi.set(qn('w:val'),'1')
    for run in p.runs:
        run.font.name='Arial'; run.font.size=Pt(11)
        rPr=run._r.get_or_add_rPr()
        rfonts=rPr.rFonts
        if rfonts is None:
            rfonts=OxmlElement('w:rFonts'); rPr.insert(0,rfonts)
        for attr in ('ascii','hAnsi','eastAsia','cs'):
            rfonts.set(qn('w:'+attr),'Arial')
        rtl=rPr.find(qn('w:rtl'))
        if rtl is None:
            rtl=OxmlElement('w:rtl'); rPr.append(rtl)
        rtl.set(qn('w:val'),'1')
        cs=rPr.find(qn('w:cs'))
        if cs is None:
            cs=OxmlElement('w:cs'); rPr.append(cs)
        cs.set(qn('w:val'),'1')

def _hyperlink(paragraph, text, url):
    part=paragraph.part; rid=part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); r=OxmlElement('w:r'); rPr=OxmlElement('w:rPr')
    color=OxmlElement('w:color'); color.set(qn('w:val'),'3D7EB9'); rPr.append(color); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rPr.append(u)
    r.append(rPr); t=OxmlElement('w:t'); t.text=text; r.append(t); h.append(r); paragraph._p.append(h)

def _clean_md(text):
    return re.sub(r'\*\*|__|^#{1,6}\s*','',text,flags=re.M).strip()

@app.post("/export/xlsx")
def export_xlsx(req:ExportReq):
    r={"answer":req.answer,"citations":req.citations,"sources":req.source_names}; wb=Workbook(); ws=wb.active; ws.title="نتيجة الراصد"; ws.sheet_view.rightToLeft=True
    navy='15445A'; green='07A869'; blue='3D7EB9'; teal='0DA9A6'; gold='C1B489'; pale='F4FAF8'; white='FFFFFF'; line='DCE8E5'
    ws.merge_cells('A1:F1'); ws['A1']='الراصد التشريعي'; ws['A1'].font=Font(name='Arial',size=20,bold=True,color=navy); ws['A1'].alignment=Alignment(horizontal='right')
    ws.merge_cells('A2:F2'); ws['A2']='دليلك إلى تحديثات الأنظمة واللوائح الحكومية السعودية'; ws['A2'].font=Font(name='Arial',size=11,bold=True,color=green); ws['A2'].alignment=Alignment(horizontal='right')
    row=4
    for label,value in [('السؤال',req.question),('المصادر المختارة','، '.join(r['sources']))]:
        ws.cell(row,1,label); ws.cell(row,1).font=Font(name='Arial',bold=True,color=white); ws.cell(row,1).fill=PatternFill('solid',fgColor=navy); ws.cell(row,1).alignment=Alignment(horizontal='right',vertical='top')
        ws.merge_cells(start_row=row,start_column=2,end_row=row,end_column=6); c=ws.cell(row,2,value); c.font=Font(name='Arial',color=navy); c.alignment=Alignment(horizontal='right',vertical='top',wrap_text=True); row+=1
    row+=1; ws.merge_cells(start_row=row,start_column=1,end_row=row,end_column=6); c=ws.cell(row,1,'نتيجة البحث والتحليل'); c.font=Font(name='Arial',size=13,bold=True,color=white); c.fill=PatternFill('solid',fgColor=green); c.alignment=Alignment(horizontal='right'); row+=1
    ws.merge_cells(start_row=row,start_column=1,end_row=row,end_column=6); c=ws.cell(row,1,_clean_md(r['answer'])); c.font=Font(name='Arial',size=11,color=navy); c.alignment=Alignment(horizontal='right',vertical='top',wrap_text=True); ws.row_dimensions[row].height=300; row+=2
    ws.merge_cells(start_row=row,start_column=1,end_row=row,end_column=6); c=ws.cell(row,1,'المصادر الرسمية المستخدمة'); c.font=Font(name='Arial',size=12,bold=True,color=white); c.fill=PatternFill('solid',fgColor=blue); c.alignment=Alignment(horizontal='right'); row+=1
    ws.append([])
    for i,cit in enumerate(r['citations'],1):
        ws.cell(row,1,i); ws.cell(row,2,cit['title']); ws.merge_cells(start_row=row,start_column=2,end_row=row,end_column=5); ws.cell(row,6,'فتح المصدر'); ws.cell(row,6).hyperlink=cit['url']; ws.cell(row,6).style='Hyperlink'
        for col in range(1,7): ws.cell(row,col).alignment=Alignment(horizontal='right',vertical='center',wrap_text=True); ws.cell(row,col).font=Font(name='Arial',color=navy); ws.cell(row,col).fill=PatternFill('solid',fgColor=pale)
        row+=1
    widths=[8,28,20,20,20,18]
    for i,w in enumerate(widths,1): ws.column_dimensions[get_column_letter(i)].width=w
    ws.freeze_panes='A4'; ws.sheet_properties.pageSetUpPr.fitToPage=True; ws.page_setup.fitToWidth=1; ws.page_margins.right=.3; ws.page_margins.left=.3
    b=io.BytesIO(); wb.save(b); b.seek(0)
    return StreamingResponse(b,media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",headers={"Content-Disposition":"attachment; filename=rasid.xlsx"})

@app.post("/export/docx")
def export_docx(req:ExportReq):
    # Word-only export. It uses the already returned search result and does NOT call the API again.
    r={"answer":req.answer,"citations":req.citations,"sources":req.source_names}
    d=Document()
    sec=d.sections[0]
    sec.top_margin=Pt(42); sec.bottom_margin=Pt(38); sec.right_margin=Pt(46); sec.left_margin=Pt(46)
    sectPr=sec._sectPr
    bidi_sec=sectPr.find(qn('w:bidi'))
    if bidi_sec is None:
        bidi_sec=OxmlElement('w:bidi'); sectPr.append(bidi_sec)
    bidi_sec.set(qn('w:val'),'1')

    NAVY='15445A'; GREEN='07A869'; BLUE='3D7EB9'; TEAL='0DA9A6'; GOLD='C1B489'; PALE='F4FAF8'; LINE='DCE8E5'; WHITE='FFFFFF'; GRAY='60777D'

    def shade(cell, color):
        tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
        if shd is None:
            shd=OxmlElement('w:shd'); tcPr.append(shd)
        shd.set(qn('w:fill'),color)

    def cell_margins(cell, top=110, start=130, bottom=110, endm=130):
        tcPr=cell._tc.get_or_add_tcPr(); mar=tcPr.first_child_found_in('w:tcMar')
        if mar is None:
            mar=OxmlElement('w:tcMar'); tcPr.append(mar)
        for tag,val in [('top',top),('start',start),('bottom',bottom),('end',endm)]:
            node=mar.find(qn('w:'+tag))
            if node is None:
                node=OxmlElement('w:'+tag); mar.append(node)
            node.set(qn('w:w'),str(val)); node.set(qn('w:type'),'dxa')

    def set_cell_text(cell, text, bold=False, color=NAVY, size=10.5):
        cell.text=''
        p=cell.paragraphs[0]
        rr=p.add_run(str(text)); rr.bold=bold; rr.font.name='Arial'; rr.font.size=Pt(size); rr.font.color.rgb=RGBColor.from_string(color)
        rr._element.rPr.rFonts.set(qn('w:cs'),'Arial')
        _rtl_paragraph(p); cell_margins(cell)
        tcPr=cell._tc.get_or_add_tcPr(); bidi=OxmlElement('w:bidiVisual'); bidi.set(qn('w:val'),'1'); tcPr.append(bidi)

    def rtl_table(table):
        tblPr=table._tbl.tblPr
        bidi=tblPr.find(qn('w:bidiVisual'))
        if bidi is None:
            bidi=OxmlElement('w:bidiVisual'); tblPr.append(bidi)
        bidi.set(qn('w:val'),'1')

    def section_title(text, color=GREEN):
        t=d.add_table(rows=1, cols=1); rtl_table(t)
        c=t.cell(0,0); shade(c,color); set_cell_text(c,text,True,WHITE,12.5)
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(1)

    def add_body(text, bold=False, color=NAVY, size=10.8):
        p=d.add_paragraph(); p.paragraph_format.space_after=Pt(5); p.paragraph_format.line_spacing=1.25
        rr=p.add_run(text); rr.bold=bold; rr.font.name='Arial'; rr.font.size=Pt(size); rr.font.color.rgb=RGBColor.from_string(color)
        rr._element.rPr.rFonts.set(qn('w:cs'),'Arial')
        _rtl_paragraph(p); return p

    def add_md_content(text):
        lines=(text or '').replace('\r','').splitlines(); i=0
        while i < len(lines):
            line=lines[i].strip()
            if not line:
                i+=1; continue
            if line.startswith('|') and line.endswith('|'):
                block=[]
                while i<len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                    block.append(lines[i].strip()); i+=1
                rows=[[x.strip() for x in z.strip('|').split('|')] for z in block]
                if len(rows)>1 and all(re.fullmatch(r'[-: ]+',x or '-') for x in rows[1]): rows.pop(1)
                if rows:
                    cols=max(len(x) for x in rows)
                    tb=d.add_table(rows=len(rows),cols=cols); rtl_table(tb); tb.style='Table Grid'; tb.autofit=True
                    for ri,row in enumerate(rows):
                        for ci in range(cols):
                            val=_clean_md(row[ci] if ci<len(row) else '')
                            set_cell_text(tb.cell(ri,ci),val,ri==0,WHITE if ri==0 else NAVY,10)
                            shade(tb.cell(ri,ci),NAVY if ri==0 else (PALE if ri%2 else WHITE))
                    d.add_paragraph()
                continue
            m=re.match(r'^(#{1,6})\s+(.*)$',line)
            if m:
                add_body(_clean_md(m.group(2)),True,NAVY,12.5 if len(m.group(1))<=2 else 11.5); i+=1; continue
            if re.match(r'^[-*]\s+',line):
                add_body('• '+_clean_md(re.sub(r'^[-*]\s+','',line)),False,NAVY,10.8); i+=1; continue
            add_body(_clean_md(line),False,NAVY,10.8); i+=1

    # True RTL in all document defaults.
    styles=d.styles
    styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(11)
    styles['Normal']._element.rPr.rFonts.set(qn('w:cs'),'Arial')
    normal_pPr=styles['Normal']._element.get_or_add_pPr(); nb=OxmlElement('w:bidi'); nb.set(qn('w:val'),'1'); normal_pPr.append(nb)
    settings=d.settings._element
    rtl=OxmlElement('w:themeFontLang'); rtl.set(qn('w:val'),'ar-SA'); settings.append(rtl)

    # Ministry-inspired clean header band.
    hdr=d.add_table(rows=1,cols=2); rtl_table(hdr); hdr.autofit=False
    hdr.columns[0].width=Inches(0.75); hdr.columns[1].width=Inches(5.9)
    for c in hdr.columns[0].cells: c.width=Inches(0.75)
    for c in hdr.columns[1].cells: c.width=Inches(5.9)
    set_cell_text(hdr.cell(0,0),'ر',True,WHITE,18); shade(hdr.cell(0,0),GREEN)
    set_cell_text(hdr.cell(0,1),'الراصد التشريعي',True,NAVY,20); shade(hdr.cell(0,1),'FFFFFF')
    p=d.add_paragraph(); rr=p.add_run('دليلك إلى تحديثات الأنظمة واللوائح الحكومية السعودية')
    rr.bold=True; rr.font.name='Arial'; rr.font.size=Pt(11.5); rr.font.color.rgb=RGBColor.from_string(GREEN); rr._element.rPr.rFonts.set(qn('w:cs'),'Arial'); _rtl_paragraph(p)
    p.paragraph_format.space_after=Pt(12)

    section_title('السؤال',NAVY)
    qtb=d.add_table(rows=1,cols=1); rtl_table(qtb); qc=qtb.cell(0,0); shade(qc,PALE); set_cell_text(qc,req.question,False,NAVY,11)
    d.add_paragraph()

    section_title('نتيجة البحث والتحليل',GREEN)
    word_answer=re.sub(r'^##\s*نتائج الراصد\s*', '', r['answer'] or '', count=1).lstrip()
    add_md_content(word_answer)

    if r.get('citations'):
        section_title('المصادر الرسمية المستخدمة',BLUE)
        for i,cit in enumerate(r['citations'],1):
            p=d.add_paragraph(); p.paragraph_format.space_after=Pt(4)
            rr=p.add_run(f'{i}. {cit["title"]} — '); rr.bold=True; rr.font.name='Arial'; rr.font.color.rgb=RGBColor.from_string(NAVY); rr.font.size=Pt(10.5); rr._element.rPr.rFonts.set(qn('w:cs'),'Arial')
            _hyperlink(p,'فتح المصدر الرسمي',cit['url']); _rtl_paragraph(p)

    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fr=footer.add_run('الراصد التشريعي  |  تحقق من النص النظامي من مصدره الرسمي قبل الاعتماد النهائي')
    fr.font.name='Arial'; fr.font.size=Pt(8); fr.font.color.rgb=RGBColor.from_string(GRAY); fr._element.rPr.rFonts.set(qn('w:cs'),'Arial')
    fpPr=footer._p.get_or_add_pPr(); bidi=OxmlElement('w:bidi'); bidi.set(qn('w:val'),'1'); fpPr.append(bidi)

    b=io.BytesIO(); d.save(b); b.seek(0)
    return StreamingResponse(b,media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",headers={"Content-Disposition":"attachment; filename=rasid.docx"})

@app.get("/",response_class=HTMLResponse)
def home():
    opts="".join(f'<label><input type="checkbox" name="src" value="{i}"> {html.escape(n)}</label>' for i,n,_ in SOURCES)
    return """<!doctype html><html dir="rtl" lang="ar"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"><title>الراصد التشريعي</title>
<style>:root{--g:#07A869;--b:#3D7EB9;--t:#0DA9A6;--n:#15445A;--gold:#C1B489;--line:#dce8e5}*{box-sizing:border-box}body{margin:0;font-family:Tahoma,Arial;color:var(--n);background:#fff}.hero{padding:42px 20px 70px;background:linear-gradient(135deg,#fff,#f3faf8);position:relative;overflow:hidden}.hero:before{content:"";position:absolute;width:520px;height:190px;border-radius:50%;background:linear-gradient(90deg,#07a86922,#3d7eb922);top:-130px;right:-70px}.inner,.shell{max-width:1100px;margin:auto}.brand{display:flex;gap:15px;align-items:center}.logo{width:58px;height:58px;border-radius:18px;background:linear-gradient(145deg,var(--g),var(--t));color:#fff;display:grid;place-items:center;font-size:28px;font-weight:bold}h1{margin:0;font-size:38px}.tag{color:var(--g);font-weight:bold;margin-top:6px;font-size:18px}.intro{margin-top:18px;color:#536c72;font-size:17px}.shell{margin-top:-35px;padding:0 18px 40px;position:relative}.card{background:#fff;border:1px solid var(--line);border-radius:22px;padding:24px;box-shadow:0 14px 38px #15445a12;margin-bottom:18px}textarea{width:100%;min-height:115px;border:1px solid #cadbd6;border-radius:15px;padding:16px;font:inherit;line-height:1.8;outline:none}textarea:focus{border-color:var(--g);box-shadow:0 0 0 4px #07a86914}.grid{display:grid;grid-template-columns:1fr 240px;gap:14px;align-items:end;margin-top:15px}.drop{position:relative}.dropbtn{width:100%;background:#fff;border:1px solid #cadbd6;color:var(--n);border-radius:13px;padding:13px;font:inherit;text-align:right}.menu{display:none;position:absolute;right:0;left:0;top:52px;background:#fff;border:1px solid var(--line);border-radius:14px;padding:8px;max-height:320px;overflow:auto;z-index:9;box-shadow:0 15px 35px #15445a25}.menu.open{display:block}.menu label{display:block;padding:9px;border-radius:8px;font-size:14px}.menu label:hover{background:#f0f8f5}.menu input{accent-color:var(--g)}button{cursor:pointer}.go{border:0;border-radius:13px;padding:14px;background:linear-gradient(135deg,var(--g),#079e75);color:#fff;font-weight:bold;font-size:16px}.go:disabled{opacity:.55}.result{line-height:1.9;white-space:normal;text-align:right;direction:rtl}.result h2{font-size:24px;margin:8px 0 14px;color:var(--n)}.result h3{font-size:17px;margin:22px 0 7px;padding:10px 12px;background:#f4faf8;border-right:4px solid var(--g);border-radius:8px;color:var(--n)}.result>div{margin:5px 0}.result strong{color:var(--n)}.citebox{margin-top:20px;padding:15px;background:#f5f8f9;border-right:4px solid var(--gold);border-radius:10px}.citebox a{color:var(--b);font-weight:bold;text-decoration:none;display:block;margin:5px 0}.actions{display:none;gap:8px;margin-top:16px}.actions button{border:1px solid #d6e5ef;background:#eef5fa;color:var(--b);border-radius:10px;padding:10px 14px;font-weight:bold}.status{margin-top:12px;color:#60777d}@media(max-width:700px){.grid{grid-template-columns:1fr}h1{font-size:30px}}</style></head>
<body><header class="hero"><div class="inner"><div class="brand"><div class="logo">ر</div><div><h1>الراصد التشريعي</h1><div class="tag">دليلك إلى تحديثات الأنظمة واللوائح الحكومية السعودية</div></div></div><div class="intro">اسأل عن نظام أو لائحة أو مادة تنظيمية، واستعرض ما تغيّر، وتحقّق من المصدر الرسمي.</div></div></header>
<main class="shell"><section class="card"><b>ماذا تريد أن تعرف؟</b><textarea id="q" placeholder="مثال: ما المواد التي تم تعديلها في اللائحة التنفيذية لنظام المنافسات والمشتريات الحكومية؟"></textarea><div class="grid"><div class="drop"><b>مصادر البحث</b><button class="dropbtn" id="dropbtn"><span id="lbl">جميع المصادر الرسمية</span> ▾</button><div id="menu" class="menu"><label><input id="all" type="checkbox" checked > جميع المصادر الرسمية</label>"""+opts+"""</div></div><button id="go" class="go" type="button">بحث وتحليل</button></div><div id="st" class="status"></div></section>
<section class="card"><b>نتائج الراصد</b><div id="out" class="result">اكتب سؤالك، وحدد المصادر التي تريد البحث فيها، ثم اضغط «بحث وتحليل».</div><div id="acts" class="actions"><button id="xlsxBtn" type="button">تصدير Excel</button><button id="docxBtn" type="button">تصدير Word</button></div></section></main>
<script>
const $=id=>document.getElementById(id);
const menu=$('menu'), all=$('all'), lbl=$('lbl'), q=$('q'), go=$('go'), st=$('st'), out=$('out'), acts=$('acts');
let last=null, lastResult=null;
function renderMd(t){let e=s=>s.replace(/[&<>]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;'}[c]));let x=e(t||'');x=x.replace(/^### (.*)$/gm,'<h3>$1</h3>').replace(/^## (.*)$/gm,'<h2>$1</h2>').replace(/^# (.*)$/gm,'<h2>$1</h2>').replace(/\*\*(.*?)\*\*/g,'<strong>$1</strong>');let lines=x.split('\\n'),o=[],tbl=[];function flush(){if(tbl.length){let rows=tbl.map(z=>z.split('|').filter(Boolean).map(v=>v.trim()));if(rows.length>1&&rows[1].every(v=>/^[-: ]+$/.test(v)))rows.splice(1,1);o.push('<div style="overflow:auto"><table style="width:100%;border-collapse:collapse;margin:12px 0">'+rows.map((r,i)=>'<tr>'+r.map(v=>'<'+(i?'td':'th')+' style="border:1px solid #dce8e5;padding:9px;text-align:right;vertical-align:top">'+v+'</'+(i?'td':'th')+'>').join('')+'</tr>').join('')+'</table></div>');tbl=[]}}for(let l of lines){if(/^\s*\|.*\|\s*$/.test(l)){tbl.push(l);continue}flush();if(l.trim())o.push('<div>'+l+'</div>')}flush();return o.join('')}
$('dropbtn').addEventListener('click',()=>menu.classList.toggle('open'));
document.addEventListener('click',e=>{if(!e.target.closest('.drop'))menu.classList.remove('open')});
document.querySelectorAll('input[name=src]').forEach(x=>x.addEventListener('change',()=>{if(x.checked)all.checked=false;upd()}));
all.addEventListener('change',()=>{if(all.checked)document.querySelectorAll('input[name=src]').forEach(x=>x.checked=false);upd()});
function upd(){let s=[...document.querySelectorAll('input[name=src]:checked')];if(!s.length){all.checked=true;lbl.textContent='جميع المصادر الرسمية'}else lbl.textContent=s.length==1?s[0].parentElement.innerText.trim():s.length+' مصادر محددة'}
function payload(){let s=[...document.querySelectorAll('input[name=src]:checked')].map(x=>x.value);return {question:q.value.trim(),sources:all.checked||!s.length?['all']:s}}
async function runSearch(){let p=payload();if(!p.question){st.textContent='اكتب سؤالك أولًا.';q.focus();return}last=p;lastResult=null;go.disabled=true;go.textContent='جارٍ البحث…';st.textContent='جارٍ البحث في المصادر الرسمية وتحليل النتائج…';acts.style.display='none';try{let r=await fetch('/analyze',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(p)});let d=await r.json();if(!r.ok)throw new Error(d.detail||'تعذر البحث');lastResult=d;out.innerHTML=renderMd(d.answer);let box=document.createElement('div');box.className='citebox';box.innerHTML='<b>المصادر الرسمية المستخدمة:</b>';(d.citations||[]).forEach((c,i)=>{let a=document.createElement('a');a.href=c.url;a.target='_blank';a.rel='noopener';a.textContent=(i+1)+'. '+c.title;box.appendChild(a)});out.appendChild(box);acts.style.display='flex';st.textContent=''}catch(e){out.textContent='تعذر تنفيذ البحث: '+e.message;st.textContent=''}finally{go.disabled=false;go.textContent='بحث وتحليل'}}
async function exp(ext){if(!last||!lastResult)return;let body={...last,answer:lastResult.answer,citations:lastResult.citations||[],source_names:lastResult.sources||[]};let r=await fetch('/export/'+ext,{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});if(!r.ok){alert('تعذر إنشاء ملف التصدير');return}let b=await r.blob(),u=URL.createObjectURL(b),a=document.createElement('a');a.href=u;a.download='نتائج_الراصد.'+ext;document.body.appendChild(a);a.click();a.remove();setTimeout(()=>URL.revokeObjectURL(u),1000)}
go.addEventListener('click',runSearch);$('xlsxBtn').addEventListener('click',()=>exp('xlsx'));$('docxBtn').addEventListener('click',()=>exp('docx'));
</script></body></html>"""
